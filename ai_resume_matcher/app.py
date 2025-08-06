from match_resume import match_resume
from rate_candidates import rate_candidates
from improve_resume import improve_resume
import streamlit as st
from generate_cover_letter import generate_cover_letter
from utils import extract_text_from_file
from io import BytesIO
from docx import Document

st.set_page_config(page_title="AI HR Assistant", layout="centered")
st.title("🤖 AI HR Assistant")


# Navigation
page = st.sidebar.radio(
    "Choose a function:",
    [
        "Resume Matching with Job Description",
        "Generate Cover Letter",
        "Candidate Evaluation",
        "Improve Resume",
    ],
)

# ========= Matching =========
if page == "Resume Matching with Job Description":
    st.header("📄 Resume Matching with Job Description")

    st.subheader("Resume")
    resume_input_type = st.radio(
        "How would you like to provide the resume?",
        ["Enter manually", "Upload a file"],
        key="resume_match",
    )
    resume_text = (
        st.text_area("Enter resume text:")
        if resume_input_type == "Enter manually"
        else ""
    )
    if resume_input_type == "Upload a file":
        resume_file = st.file_uploader(
            "Resume file (PDF/TXT/DOCX)",
            type=["pdf", "txt", "docx"],
            key="resume_file_match",
        )
        if resume_file:
            resume_text = extract_text_from_file(resume_file)

    st.subheader("Job Description")
    job_input_type = st.radio(
        "How would you like to provide the job description?",
        ["Enter manually", "Upload a file"],
        key="job_match",
    )
    job_text = (
        st.text_area("Enter job description text:")
        if job_input_type == "Enter manually"
        else ""
    )
    if job_input_type == "Upload a file":
        job_file = st.file_uploader(
            "Job description file (PDF/TXT/DOCX)",
            type=["pdf", "txt", "docx"],
            key="job_file_match",
        )
        if job_file:
            job_text = extract_text_from_file(job_file)

    if st.button("🔍 Match") and resume_text and job_text:
        result = match_resume(resume_text, job_text)
        st.subheader("🔎 Matching Result:")
        st.write(result)

# ========= Cover letter =========
elif page == "Generate Cover Letter":
    st.header("✉️ Generate Cover Letter")

    st.subheader("Resume")
    resume_input_type = st.radio(
        "How would you like to provide the resume?",
        ["Enter manually", "Upload a file"],
        key="resume_cover",
    )
    resume_text = ""
    if resume_input_type == "Enter manually":
        resume_text = st.text_area("Enter resume text:")
    else:
        resume_file = st.file_uploader(
            "Upload resume file (PDF/TXT/DOCX)",
            type=["pdf", "txt", "docx"],
            key="resume_file_cover",
        )
        if resume_file:
            resume_text = extract_text_from_file(resume_file)

    st.subheader("Job Description")
    job_input_type = st.radio(
        "How would you like to provide the job description?",
        ["Enter manually", "Upload a file"],
        key="job_cover",
    )
    job_text = ""
    if job_input_type == "Enter manually":
        job_text = st.text_area("Enter job description text:")
    else:
        job_file = st.file_uploader(
            "Upload job description file (PDF/TXT/DOCX)",
            type=["pdf", "txt", "docx"],
            key="job_file_cover",
        )
        if job_file:
            job_text = extract_text_from_file(job_file)

    if st.button("📨 Generate Letter") and resume_text and job_text:
        cover_letter = generate_cover_letter(resume_text, job_text)
        st.subheader("📄 Cover Letter:")
        st.text_area("Copy the text below:", value=cover_letter, height=300)

        # Download txt
        txt_bytes = BytesIO(cover_letter.encode("utf-8"))
        st.download_button(
            "⬇️ Download as .txt",
            data=txt_bytes,
            file_name="cover_letter.txt",
            mime="text/plain",
        )

        # Download DOCX
        docx_file = BytesIO()
        doc = Document()
        doc.add_paragraph(cover_letter)
        doc.save(docx_file)
        docx_file.seek(0)
        st.download_button(
            "⬇️ Download as .docx",
            data=docx_file,
            file_name="cover_letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

elif page == "Candidate Evaluation":
    st.header("👥 Candidate Evaluation")

    # Выбор способа ввода описания вакансии
    job_input_method = st.radio("How would you like to provide the job description?", ["📄 Upload a file", "📝 Enter manually"])

    # Ввод описания вакансии
    job_text = ""
    if job_input_method == "📄 Upload a file":
        job_file = st.file_uploader("Upload job description", type=["pdf", "txt", "docx"])
        if job_file:
            job_text = extract_text_from_file(job_file)
    else:
        job_text = st.text_area("Enter the job description manually", height=300)

    # Загрузка резюме
    resume_files = st.file_uploader(
        "Upload candidate resumes",
        type=["pdf", "txt", "docx"],
        accept_multiple_files=True,
    )

    if st.button("📊 Evaluate Candidates"):
        if not job_text.strip() or not resume_files:
            st.error("❌ Please provide a job description and at least one resume.")
            st.stop()

        # Обработка резюме
        resume_texts = []
        resume_names = []
        for resume_file in resume_files:
            text = extract_text_from_file(resume_file)
            resume_texts.append(text)
            resume_names.append(resume_file.name)

        try:
            result = rate_candidates(resume_texts, job_text)

            st.subheader("🏆 Candidate Ranking:")
            st.text_area("Result", result, height=500)

            # Скачать как .txt
            st.download_button(
                label="📥 Download as .txt",
                data=result,
                file_name="candidate_ranking.txt",
                mime="text/plain",
            )

            # Скачать как .docx
            from docx import Document
            from io import BytesIO

            doc = Document()
            doc.add_heading("🏆 Candidate Ranking", level=1)

            for line in result.split("\n"):
                doc.add_paragraph(line)

            docx_io = BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)

            st.download_button(
                label="📄 Download as .docx",
                data=docx_io,
                file_name="candidate_ranking.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        except Exception as e:
            st.error(f"Error during ranking: {e}")

# ========= Resume Improvement =========
elif page == "Improve Resume":
    st.header("🛠️ Improve Resume")

    resume_input_method = st.radio(
        "Choose resume input method:",
        ["Enter manually", "Upload file (.pdf, .txt, .docx)"],
    )
    resume_text = ""
    if resume_input_method == "Enter manually":
        resume_text = st.text_area("Enter your resume:", height=300)
    else:
        resume_file = st.file_uploader(
            "Upload resume file", type=["pdf", "txt", "docx"]
        )
        if resume_file:
            resume_text = extract_text_from_file(resume_file)

    job_input_method = st.radio(
        "Choose job description input method:",
        ["Enter manually", "Upload file (.pdf, .txt, .docx)"],
    )
    job_description = ""
    if job_input_method == "Enter manually":
        job_description = st.text_area("Enter job description:", height=250)
    else:
        job_file = st.file_uploader(
            "Upload job description file", type=["pdf", "txt", "docx"]
        )
        if job_file:
            job_description = extract_text_from_file(job_file)

    if resume_text and job_description and st.button("✨ Improve"):
        with st.spinner("Processing..."):
            improved = improve_resume(resume_text, job_description)

        st.subheader("💡 Improved Resume:")
        st.text_area("Result:", value=improved, height=300)

        st.download_button(
            "📄 Download as TXT", improved, file_name="improved_resume.txt"
        )

        from docx import Document
        import tempfile, os

        docx_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc = Document()
        for line in improved.split("\n"):
            doc.add_paragraph(line)
        doc.save(docx_file.name)
        with open(docx_file.name, "rb") as f:
            st.download_button(
                "📄 Download as DOCX", f.read(), file_name="improved_resume.docx"
            )
        os.remove(docx_file.name)
